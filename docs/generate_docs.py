#!/usr/bin/env python3.12
"""
Washermann Financial & Pricing Module — Comprehensive Documentation Generator
Produces a fully formatted DOCX document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colour palette ───────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB
BRAND_DARK   = RGBColor(0x11, 0x18, 0x27)   # #111827
BRAND_SLATE  = RGBColor(0x6B, 0x72, 0x80)   # #6B7280
BRAND_GREEN  = RGBColor(0x05, 0x96, 0x69)   # #059669
BRAND_AMBER  = RGBColor(0xD9, 0x77, 0x06)   # #D97706
BRAND_RED    = RGBColor(0xDC, 0x26, 0x26)   # #DC2626
TABLE_HEADER = RGBColor(0x1E, 0x3A, 0x5F)   # #1E3A5F dark navy
TABLE_ALT    = RGBColor(0xF0, 0xF4, 0xFF)   # very light blue
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY   = RGBColor(0xF3, 0xF4, 0xF6)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_bookmark(para, bookmark_id: int, name: str):
    """Add a Word bookmark to a paragraph (used for cross-references)."""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'),   str(bookmark_id))
    start.set(qn('w:name'), name)
    para._p.append(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    para._p.append(end)


class DocBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._bm_id = 0

    # ─── Style setup ──────────────────────────────────────────────────────────

    def _setup_styles(self):
        doc = self.doc
        section = doc.sections[0]
        section.page_width   = Inches(8.5)
        section.page_height  = Inches(11)
        section.left_margin  = Inches(1.1)
        section.right_margin = Inches(1.1)
        section.top_margin   = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        # Normal
        normal = doc.styles['Normal']
        normal.font.name  = 'Calibri'
        normal.font.size  = Pt(10.5)
        normal.font.color.rgb = BRAND_DARK

        # Heading 1
        h1 = doc.styles['Heading 1']
        h1.font.name   = 'Calibri'
        h1.font.size   = Pt(22)
        h1.font.bold   = True
        h1.font.color.rgb = BRAND_BLUE
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after  = Pt(8)

        # Heading 2
        h2 = doc.styles['Heading 2']
        h2.font.name  = 'Calibri'
        h2.font.size  = Pt(15)
        h2.font.bold  = True
        h2.font.color.rgb = BRAND_DARK
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after  = Pt(6)

        # Heading 3
        h3 = doc.styles['Heading 3']
        h3.font.name  = 'Calibri'
        h3.font.size  = Pt(12)
        h3.font.bold  = True
        h3.font.color.rgb = BRAND_BLUE
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after  = Pt(4)

        # Heading 4
        h4 = doc.styles['Heading 4']
        h4.font.name  = 'Calibri'
        h4.font.size  = Pt(10.5)
        h4.font.bold  = True
        h4.font.italic = False
        h4.font.color.rgb = BRAND_SLATE
        h4.paragraph_format.space_before = Pt(8)
        h4.paragraph_format.space_after  = Pt(2)

    # ─── Paragraph builders ───────────────────────────────────────────────────

    def h1(self, text, bm=None):
        p = self.doc.add_heading(text, level=1)
        if bm:
            self._bm_id += 1
            add_bookmark(p, self._bm_id, bm)
        return p

    def h2(self, text, bm=None):
        p = self.doc.add_heading(text, level=2)
        if bm:
            self._bm_id += 1
            add_bookmark(p, self._bm_id, bm)
        return p

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    def h4(self, text):
        return self.doc.add_heading(text, level=4)

    def body(self, text, bold=False, italic=False, color=None, size=None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(6)
        return p

    def lead(self, text):
        """Larger intro paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11.5)
        run.font.color.rgb = BRAND_SLATE
        p.paragraph_format.space_after = Pt(10)
        return p

    def bullet(self, text, level=0, bold_prefix=None):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent   = Inches(0.3 + 0.25 * level)
        p.paragraph_format.space_after   = Pt(3)
        if bold_prefix:
            run = p.add_run(bold_prefix + ': ')
            run.bold = True
            run.font.color.rgb = BRAND_DARK
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def numbered(self, text, bold_prefix=None):
        p = self.doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            run = p.add_run(bold_prefix + ': ')
            run.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    def note(self, text, label='NOTE', color=None):
        """Styled callout box (simulated with a single-cell table)."""
        color = color or BRAND_AMBER
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        cell = tbl.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run(f'ℹ  {label}: ')
        run.bold = True
        run.font.color.rgb = color
        run2 = p.add_run(text)
        run2.font.color.rgb = BRAND_DARK
        run2.font.size = Pt(10)
        set_cell_bg(cell, 'FFF8E7')
        self.doc.add_paragraph()  # breathing room

    def code(self, text):
        """Monospace code block."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(2)
        return p

    def divider(self):
        self.doc.add_paragraph()

    def page_break(self):
        self.doc.add_page_break()

    # ─── Table builder ────────────────────────────────────────────────────────

    def table(self, headers, rows, col_widths=None):
        """Render a formatted table with dark header row and alternating row shading."""
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            set_cell_bg(cell, '1E3A5F')
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

        # Data rows
        for ri, row in enumerate(rows):
            row_cells = tbl.rows[ri + 1].cells
            bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
            for ci, val in enumerate(row):
                cell = row_cells[ci]
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                if isinstance(val, tuple):
                    run = p.add_run(val[0])
                    run.bold   = val[1] if len(val) > 1 else False
                    run.italic = val[2] if len(val) > 2 else False
                else:
                    run = p.add_run(str(val))
                run.font.size = Pt(9.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)

        # Column widths
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in tbl.rows:
                    row.cells[i].width = Inches(w)

        self.doc.add_paragraph()  # space after table
        return tbl

    # ─── Save ─────────────────────────────────────────────────────────────────

    def save(self, path):
        self.doc.save(path)
        print(f'Saved → {path}')


# ─── Build the document ───────────────────────────────────────────────────────

def build(path: str):
    b = DocBuilder()
    doc = b.doc

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('WASHERMANN')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(36)
    run.font.bold  = True
    run.font.color.rgb = BRAND_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Financial & Pricing Module')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(22)
    run.font.bold  = False
    run.font.color.rgb = BRAND_DARK

    doc.add_paragraph()
    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Comprehensive Technical & Business Documentation')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(13)
    run.font.color.rgb = BRAND_SLATE

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 1.0  ·  Confidential  ·  May 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_SLATE

    doc.add_paragraph()
    doc.add_paragraph()

    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = audience.add_run('Audience: Engineering  ·  Product  ·  Finance  ·  Investors  ·  QA')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BRAND_SLATE

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('Table of Contents')

    toc_entries = [
        ('1.', 'Executive Summary', '3'),
        ('2.', 'Platform Overview & Architecture Philosophy', '4'),
        ('3.', 'WashPoints — The Internal Currency', '5'),
        ('4.', 'Conversion Rate System', '7'),
        ('5.', 'Pricing Architecture', '8'),
        ('   5.1', 'Platform Price List (Catalogue)', '8'),
        ('   5.2', 'Pricing Engine', '9'),
        ('   5.3', 'Fee Structure', '11'),
        ('   5.4', 'All-In Pricing Model Endpoint', '12'),
        ('6.', 'Special Pricing Packages', '13'),
        ('7.', 'Vendor Price Intelligence Engine', '15'),
        ('8.', 'Platform Configuration', '17'),
        ('9.', 'Order Financial Flow (End-to-End)', '19'),
        ('10.', 'Escrow System', '21'),
        ('11.', 'Customer Wallet & Fiat Cost Basis Tracking', '22'),
        ('12.', 'Vendor Earnings System', '24'),
        ('13.', 'Rep Pseudo-Wallet & Bonus Tiers', '25'),
        ('14.', 'Payout System', '27'),
        ('15.', 'Gift Cards & Vaults', '28'),
        ('16.', 'Audit Trail & Ledger Immutability', '30'),
        ('17.', 'Profile Completion Gate', '31'),
        ('18.', 'API Endpoint Reference', '32'),
        ('19.', 'QA Test Scenarios', '35'),
        ('20.', 'Glossary', '38'),
    ]

    for num, title_text, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run1 = p.add_run(f'{num}  {title_text}')
        run1.font.size = Pt(10.5)
        if not title_text.startswith(' ') and num.count('.') == 1 and '.' in num:
            run1.bold = True
        dots = '.' * max(2, 60 - len(num) - len(title_text))
        run2 = p.add_run(f'  {dots}  {page}')
        run2.font.size = Pt(10)
        run2.font.color.rgb = BRAND_SLATE

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('1.  Executive Summary')

    b.lead(
        'Washermann is a laundry-as-a-service marketplace connecting customers to '
        'vetted laundry vendors and delivery representatives ("Reps") in Nigeria. '
        'This document covers the complete financial and pricing architecture that '
        'powers every transaction on the platform.'
    )

    b.body(
        'The financial system is designed around four core principles:'
    )
    b.bullet('Predictability — customers always know what they will pay before confirming an order.')
    b.bullet('Fairness — prices are informed by real vendor market data, not arbitrary markups.')
    b.bullet('Flexibility — promotional pricing, package bundles, and audience-targeted offers can be deployed without engineering changes.')
    b.bullet('Auditability — every penny movement leaves an immutable, timestamped record in a double-entry ledger system.')

    b.divider()
    b.body('Key Financial Metrics (illustrative defaults):', bold=True)
    b.table(
        headers=['Metric', 'Default Value', 'Configurable?'],
        rows=[
            ('WashPoints per ₦1 (conversion rate)', 'Admin-set (e.g. 10 WP/₦1)', 'Yes — live'),
            ('Service charge on order subtotal', '5%', 'Yes — live'),
            ('VAT', '0% (can be enabled)', 'Yes — live'),
            ('Rep share of order total', '15%', 'Yes — live'),
            ('Vendor payout rate', '₦9 per WashPoint', 'Yes — live'),
            ('Platform price offset over vendor median', '25%', 'Yes — live'),
            ('Price suggestion percentile (P70 default)', '70th percentile', 'Yes — live'),
            ('Order auto-complete (after delivery)', '24 hours', 'Yes — live'),
        ],
        col_widths=[2.8, 2.2, 1.3]
    )

    b.note(
        'All percentage values can be changed live by an admin without a deployment. '
        'Changes take effect on the next order placed — existing orders are never retroactively modified.',
        label='INVESTOR NOTE', color=BRAND_GREEN
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — PLATFORM OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('2.  Platform Overview & Architecture Philosophy')

    b.h2('2.1  Who the Participants Are')
    b.table(
        headers=['Participant', 'Role', 'Financial Relationship'],
        rows=[
            ('Customer', 'Places laundry orders; pays in WashPoints', 'Buys WP with Naira via Paystack; WP debited on order placement'),
            ('Vendor', 'Laundry business that processes garments', 'Receives WP earnings per completed order; cashes out to Naira bank account'),
            ('Rep (Delivery)', 'Collects from customer, delivers to vendor and back', 'Earns WP per delivery logged in a pseudo-wallet; paid out as Naira salary each cycle'),
            ('Admin / Finance', 'Manages platform config, approves payouts, monitors intelligence', 'Controls all fee structures, conversion rates, and price catalogue'),
        ],
        col_widths=[1.2, 2.4, 2.7]
    )

    b.h2('2.2  Why a Point-Based Currency?')
    b.body(
        'WashPoints (WP) are the internal unit of account for every transaction on '
        'Washermann. This design decision provides several advantages:'
    )
    b.bullet('Rate insulation', bold_prefix='')
    b.body('The platform can adjust the Naira↔WP conversion rate without touching price catalogue entries. '
           'A bag that costs 1,000 WP stays at 1,000 WP even if the Naira value changes.', italic=True)
    b.bullet('Loyalty potential', bold_prefix='')
    b.body('Points create natural stickiness. Customers who have accumulated WP are more likely to return.', italic=True)
    b.bullet('Multi-currency readiness', bold_prefix='')
    b.body('The internal WP layer can be priced against USD, GHS, or any other currency by simply adding a new ConversionRate row.', italic=True)
    b.bullet('Fraud surface reduction', bold_prefix='')
    b.body('WP cannot be withdrawn directly. They can only be spent on orders, gift cards, or package bundles, limiting exposure to financial fraud.', italic=True)

    b.h2('2.3  Architecture at a Glance')
    b.body('The system is built on NestJS 11 + TypeORM 0.3 + PostgreSQL. '
           'Every financial operation runs inside a database transaction to guarantee atomicity. '
           'No financial record is ever updated or deleted — only new records are appended.')

    b.table(
        headers=['Layer', 'Responsibility'],
        rows=[
            ('PricingEngine (stateless)',    'Pure calculation — inputs in, PricingResult out. Zero database access.'),
            ('PricingService',               'Loads config from DB, calls engine, orchestrates getClientConfig.'),
            ('PricingIntelligenceService',   'Aggregates vendor prices, computes stats, suggests & applies platform prices.'),
            ('PricingPackagesService',       'CRUD for promotional packages; audience-matching for customer-facing display.'),
            ('PlatformConfigService',        'Single-row config store + immutable price list CRUD.'),
            ('OrdersService',               'Authoritative order placement; wallet debit; escrow creation; earnings split.'),
            ('Ledger entities',              'Append-only audit trail for customer wallets, vendor earnings, rep pseudo-wallets.'),
        ],
        col_widths=[2.4, 3.9]
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — WASHPOINTS CURRENCY
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('3.  WashPoints — The Internal Currency')

    b.lead(
        'WashPoints (WP) are the exclusive unit of exchange inside Washermann. '
        'Every order, refund, gift card, and earnings payout is denominated in WP. '
        'Customers buy WP with Naira; vendors earn WP and cash out to Naira.'
    )

    b.h2('3.1  Acquiring WashPoints (Customer)')
    b.body('Customers top up their WashPoint balance by paying via Paystack:')
    b.numbered('Customer selects a WP amount to purchase on the app.')
    b.numbered('The system locks in the current conversion rate (e.g. 10 WP / ₦1).')
    b.numbered('A Paystack transaction record is created with the rate snapshot — this rate is frozen at initiation, not at payment completion, preventing rate-manipulation exploits.')
    b.numbered('Customer completes payment (card, bank transfer, USSD, etc.).')
    b.numbered('Paystack fires a webhook; the system verifies it independently.')
    b.numbered('WP are credited to the customer\'s Wallet and a LedgerEntry is written.')
    b.body('')
    b.note(
        'The conversion rate is snapshotted at the moment the customer clicks "Buy WashPoints", '
        'not when the Paystack webhook arrives. This means a rate change mid-transaction does '
        'not give the customer more or fewer points than they bargained for.',
        label='IMPORTANT'
    )

    b.h2('3.2  Spending WashPoints')
    b.body('WP are debited from the customer\'s wallet at the moment an order is placed. The amount is determined by the authoritative server-side pricing calculation, never by a client-supplied figure.')

    b.h2('3.3  Earning WashPoints (Vendors & Reps)')
    b.body('Vendors earn WP for each order they complete. These WP accumulate in a VendorEarningsWallet and can be requested as a Naira payout at any time.')
    b.body('Reps earn WP in a pseudo-wallet (not directly visible to them). Ops converts their accumulated cycle WP to a Naira salary each bonus cycle.')

    b.h2('3.4  Fiat Cost-Basis Tracking (WACB Method)')
    b.body(
        'The wallet entity stores both the WP balance and a parallel fiatBalanceKobo field. '
        'This tracks the real Naira value the customer has paid for their remaining WP, '
        'using Weighted Average Cost Basis (WACB) — the same method accountants use for '
        'stock portfolios.'
    )

    b.h3('How it works step by step:')
    b.numbered('Top-up: Customer pays ₦10,000 and receives 100,000 WP at 10 WP/₦1. '
               'fiatBalanceKobo increases by 1,000,000 kobo (= ₦10,000). WP balance = 100,000 WP.')
    b.numbered('Top-up again: Customer pays ₦5,000 and receives 45,000 WP at 9 WP/₦1 (rate changed). '
               'fiatBalanceKobo += 500,000 kobo. WP balance = 145,000 WP. '
               'Effective cost basis = ₦15,000 / 145,000 WP ≈ ₦0.103/WP.')
    b.numbered('Order placed for 50,000 WP. Proportion debited = 50,000 / 145,000 ≈ 34.5%. '
               'fiatBalanceKobo decreases by 34.5% × 1,500,000 = 517,241 kobo. '
               'Remaining balance = 95,000 WP, fiatBalance = 982,759 kobo (≈ ₦9,828).')

    b.body('This design enables:', bold=True)
    b.bullet('Tax reporting — the system can report how much Naira a customer actually spent')
    b.bullet('Refund accuracy — refunds restore proportional Naira value, not an arbitrary WP amount')
    b.bullet('Fraud detection — large discrepancies between WP balance and fiat balance flag anomalies')

    b.note(
        'WACB is simpler than FIFO lot tracking (which would require a separate lots table) '
        'and mathematically correct for tax purposes. A single column on Wallet is all that is needed.',
        label='TECHNICAL NOTE', color=BRAND_GREEN
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — CONVERSION RATE SYSTEM
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('4.  Conversion Rate System')

    b.lead(
        'The ConversionRate table stores the official Naira-to-WashPoint exchange rate. '
        'Multiple rates can coexist; the "active" rate is the most recent one whose '
        'effectiveFrom is in the past.'
    )

    b.h2('4.1  How Rates Are Applied')
    b.table(
        headers=['Event', 'Rate Used', 'Behaviour'],
        rows=[
            ('Customer buys WP via Paystack',     'Rate at initiation',   'Snapshot locked when transaction is created — not when webhook fires'),
            ('Order pricing display (client)',     'Current active rate',  'Used to convert WP totals to Naira for display only'),
            ('Order pricing (server, at placement)', 'Current active rate', 'Snapshot stored in pricingSnapshot on the Order record'),
            ('Vendor earnings display',            'Rate at order time',   'nairaSnapshot stored in VendorLedgerEntry for historical reference'),
            ('Vendor payout request',              'payoutRateNairaPerWP', 'Separate from the buy-rate; admin-configurable; snapshotted at request time'),
        ],
        col_widths=[2.2, 1.6, 2.5]
    )

    b.h2('4.2  Buy Rate vs Payout Rate')
    b.body(
        'There are two distinct rates in the system:'
    )
    b.bullet('Buy rate (pointsPerUnit on ConversionRate)', bold_prefix='')
    b.body('How many WashPoints the customer gets per ₦1 spent. '
           'Higher = more generous to customers = lower effective price.', italic=True)
    b.bullet('Payout rate (payoutRateNairaPerWP on PlatformConfig)', bold_prefix='')
    b.body('How many Naira a vendor or rep receives per WashPoint. '
           'This is the cost-of-goods-sold rate for the platform.', italic=True)

    b.note(
        'The spread between buy rate and payout rate is a key lever for platform profitability. '
        'Example: If customers buy at 10 WP/₦1 (₦0.10/WP) but vendors are paid at ₦0.09/WP, '
        'the platform retains ₦0.01 per WP on top of service charge. '
        'This spread, combined with the service charge, is the platform\'s primary revenue.',
        label='INVESTOR NOTE', color=BRAND_GREEN
    )

    b.h2('4.3  Multi-Currency Readiness')
    b.body(
        'The ConversionRate entity has a currency field (default: "NGN"). '
        'Adding a second row with currency = "USD" immediately enables dollar-denominated '
        'WP purchases without any code change. The active rate is always the most recent '
        'approved entry per currency.'
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — PRICING ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('5.  Pricing Architecture')

    b.lead(
        'Washermann pricing is built in three layers: a static price catalogue '
        '(what things cost in WP), a stateless calculation engine (what the customer '
        'actually pays), and an intelligent suggestion system (how prices stay competitive).'
    )

    # 5.1 Price List
    b.h2('5.1  Platform Price List (Catalogue)')

    b.body(
        'The PlatformPriceList table is an append-only catalogue of WashPoint prices. '
        'No entry is ever updated or deleted — price changes are implemented by adding '
        'a new row with a future effectiveFrom date. The system always uses the most '
        'recent approved row that has become effective.'
    )

    b.h3('Price entry types:')
    b.table(
        headers=['priceType', 'Key Fields', 'Example'],
        rows=[
            ('bag',          'serviceType + bagSize',  '"wash_fold medium bag" = 1,200 WP'),
            ('special_item', 'itemType',               '"suit" = 800 WP, "agbada" = 1,500 WP, "duvet" = 600 WP'),
            ('ironing',      '(none — single global price)',  '"per garment ironing" = 80 WP'),
        ],
        col_widths=[1.2, 2.3, 2.8]
    )

    b.h3('Bag size matrix (illustrative):')
    b.table(
        headers=['Bag Size', 'Wash & Fold (WP)', 'Wash & Iron (WP)', 'Approximate Naira (at 10 WP/₦1)'],
        rows=[
            ('Small',  '800',   '1,100',  '₦80 – ₦110'),
            ('Medium', '1,200', '1,600',  '₦120 – ₦160'),
            ('Large',  '1,800', '2,400',  '₦180 – ₦240'),
            ('XL',     '2,500', '3,200',  '₦250 – ₦320'),
        ],
        col_widths=[1.0, 1.7, 1.7, 1.9]
    )
    b.body('All values above are illustrative. Actual prices are configured by admin.', italic=True, color=BRAND_SLATE)

    b.h3('Append-only rationale:')
    b.body(
        'By never modifying existing price entries, the system preserves a complete '
        'price history. When auditing an old order, the exact prices in effect at that '
        'time can always be reconstructed. This also means pricing disputes can be '
        'resolved definitively.'
    )

    b.divider()

    # 5.2 Pricing Engine
    b.h2('5.2  Pricing Engine')

    b.body(
        'The PricingEngine is a standalone, stateless TypeScript class. It has no '
        'database access and no NestJS dependencies. It accepts a PricingInputs object '
        'and a PricingConfig object, and returns a deterministic PricingResult.'
    )

    b.h3('Calculation Steps (in order):')
    b.table(
        headers=['Step', 'What Happens', 'Applied To'],
        rows=[
            ('1. Bag price',      'Base price for service type + bag size',                     'Every order'),
            ('2. Special items',  'Per-item price × quantity for each requested item type',     'If customer adds special items'),
            ('3. Ironing add-on', 'Per-garment iron price × ironing count',                    'wash_iron orders only'),
            ('4. Subtotal',       'Sum of steps 1–3',                                           'Derived'),
            ('5. Service charge', 'subtotal × serviceChargePercent ÷ 100',                     'Every order'),
            ('6. VAT',            '(subtotal + serviceCharge) × vatPercent ÷ 100',             'When vatPercent > 0'),
            ('7. Transport fee',  'Flat fee from area config (transportFeeWP)',                 'Every order'),
            ('8. Total',          'subtotal + serviceCharge + VAT + transport',                 'Final charge to customer'),
            ('9. Naira equiv.',   'totalWP ÷ pointsPerUnit (display only, never stored)',       'For UI display'),
        ],
        col_widths=[1.4, 3.1, 1.8]
    )

    b.h3('Full PricingResult object:')
    b.code(
        'lineItems[]         — itemised breakdown (label, category, unitPriceWP, qty, subtotalWP)\n'
        'subtotalWP          — sum before fees\n'
        'serviceChargeWP     — calculated service charge\n'
        'vatWP               — calculated VAT (0 when disabled)\n'
        'transportWP         — area transport fee\n'
        'totalWP             — final amount charged to customer\n'
        'nairaEquivalent     — display-only Naira value\n'
        'conversionRateId    — UUID of rate used\n'
        'conversionRateSnapshot — pointsPerUnit at calculation time\n'
        'calculatedAt        — ISO 8601 timestamp'
    )

    b.h3('Why the engine is stateless:')
    b.body(
        'The engine performs no I/O. This means it can be unit-tested exhaustively '
        'with zero mocking, runs identically on server and (if needed) on a client device, '
        'and has a provably deterministic output for any given inputs. '
        'The server runs it authoritatively at order placement; the client can run the '
        'same logic locally for real-time cart previews without a network call per update.'
    )

    b.note(
        'The client is NEVER trusted to provide a final price. The server re-runs the full '
        'calculation at order placement. The client-side calculation is for UX speed only.',
        label='SECURITY NOTE', color=BRAND_RED
    )

    b.divider()

    # 5.3 Fee Structure
    b.h2('5.3  Fee Structure')

    b.body('Every order passes through several fee layers. Here is what each fee means and who it benefits:')

    b.table(
        headers=['Fee', 'Default', 'Applied On', 'Who Benefits', 'Configurable?'],
        rows=[
            ('Service charge', '5%',   'Order subtotal (items only)',              'Platform revenue',          'Yes'),
            ('VAT',            '0%',   'Subtotal + service charge',                'Government / compliance',   'Yes'),
            ('Transport fee',  'Varies by area', 'Flat per order',                 'Covers Rep logistics cost', 'Per area'),
            ('Rep share',      '15%',  'Total order value (informational)',        'Rep earnings pool',         'Yes'),
            ('Platform spread','Implicit', 'Buy rate vs payout rate difference',   'Platform margin',           'Via rate config'),
        ],
        col_widths=[1.3, 1.1, 1.8, 1.7, 1.0]
    )

    b.h3('Worked Example:')
    b.body('Customer orders a Medium Wash & Fold bag with one suit. Platform: 5% service charge, 0% VAT, ₦150 transport.')
    b.table(
        headers=['Line Item', 'WP', 'Naira (at 10 WP/₦1)'],
        rows=[
            ('Medium Bag — Wash & Fold',  '1,200',   '₦120.00'),
            ('Suit',                       '800',     '₦80.00'),
            ('Subtotal',                   '2,000',   '₦200.00'),
            ('Service charge (5%)',        '100',     '₦10.00'),
            ('Transport',                  '1,500',   '₦150.00'),
            (('Total', True),              ('3,600', True),  ('₦360.00', True)),
        ],
        col_widths=[2.8, 1.2, 1.5]
    )

    b.divider()

    # 5.4 All-In Pricing Model Endpoint
    b.h2('5.4  All-In Pricing Model Endpoint  (GET /pricing/model/:areaId)')

    b.body(
        'This endpoint returns the complete pricing model for a given delivery area. '
        'Its primary purpose is to give the mobile/web client everything it needs to '
        'show customers the final, all-inclusive price without a separate API call per '
        'cart change.'
    )

    b.h3('What "all-in" means:')
    b.body(
        'Every item price is returned in two forms. '
        'The rawWP is the base platform price. The totalWP is the raw price after the '
        'service charge and VAT multiplier has been applied. Transport is excluded from '
        'item prices because it is a per-order flat fee, not a per-item charge.'
    )
    b.code(
        '// Example response structure (simplified)\n'
        '{\n'
        '  bagPrices: {\n'
        '    wash_fold: {\n'
        '      small:  { rawWP: 800,   totalWP: 840 },   // 5% SC applied\n'
        '      medium: { rawWP: 1200,  totalWP: 1260 },\n'
        '      ...\n'
        '    }\n'
        '  },\n'
        '  specialItemPrices: {\n'
        '    suit:   { rawWP: 800, totalWP: 840 },\n'
        '    agbada: { rawWP: 1500, totalWP: 1575 },\n'
        '  },\n'
        '  ironing: { rawWP: 80, totalWP: 84 },\n'
        '  transportFeeWP: 1500,\n'
        '  fees: {\n'
        '    serviceChargePercent: 5,\n'
        '    vatPercent: 0,\n'
        '    repSharePercent: 15,\n'
        '    feeMultiplier: 1.05,\n'
        '    effectivePercentage: 5.0\n'
        '  },\n'
        '  conversion: { conversionRateId: "...", pointsPerUnit: 10, nairaPerWP: 0.1 },\n'
        '  cachedAt: "2026-05-08T10:30:00.000Z"\n'
        '}'
    )

    b.body(
        'The feeMultiplier field lets the client apply the same math to any custom item '
        'not in the list: customPrice × feeMultiplier + transportFeeWP = total customer pays.'
    )

    b.note(
        'Clients should cache this response for approximately 5 minutes. '
        'The server will re-calculate authoritatively at order placement regardless of '
        'what the client cached.',
        label='IMPLEMENTATION NOTE'
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — SPECIAL PRICING PACKAGES
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('6.  Special Pricing Packages')

    b.lead(
        'Pricing Packages are promotional or niche bundles that exist outside the '
        'standard bag-size pricing system. They have a fixed WP price agreed at creation, '
        'a structured list of what is included, and powerful audience targeting rules.'
    )

    b.h2('6.1  Why Packages Exist')
    b.body(
        'The standard bag-size model works well for general laundry but cannot express '
        'offers like "15 baby garments for ₦X" or "5 corporate shirts pressed for ₦Y". '
        'Packages solve this by letting the admin define any arbitrary bundle with a '
        'fixed price, without needing to know how many bags it would equate to.'
    )

    b.h3('Example packages:')
    b.table(
        headers=['Package Name', 'What\'s Included', 'Fixed Price (WP)', 'Target Audience'],
        rows=[
            ('Baby Bundle',    '15 baby garments, gentle wash & fold',        '8,000 WP',  'New parents; users with newborns'),
            ('Corporate Pack', '5 shirts + 3 trousers, wash & press',         '12,000 WP', 'Corporate company members'),
            ('Duvet Season',   '2 duvets fully washed (limited time)',         '7,000 WP',  'All users (seasonal promo)'),
            ('VIP Welcome',    '1 full-service order for new customers',       '5,000 WP',  'Users with 0 completed orders'),
            ('Loyalty Reward', '10% off any bag size (redeemed as package)',   '6,500 WP',  'Users with 10+ completed orders'),
        ],
        col_widths=[1.4, 2.5, 1.3, 1.8]
    )

    b.h2('6.2  Package Anatomy')
    b.h3('Core fields:')
    b.bullet('name — what the customer sees (e.g. "Baby Bundle")')
    b.bullet('description — marketing copy describing the package')
    b.bullet('imageUrl — artwork for the card in the app UI')
    b.bullet('priceWP — the fixed WP price; not calculated from items')
    b.bullet('criteria — structured list of what garments are included (for display; not enforced at the service layer in this version)')
    b.bullet('audience — targeting rules (who can see this package)')
    b.bullet('isActive — on/off toggle without deletion')
    b.bullet('displayOrder — controls sort order in the app (lower = shown first)')
    b.bullet('validFrom / validUntil — time-bounded promotional offers')
    b.bullet('maxUsesPerUser — per-user usage cap (null = unlimited)')

    b.h2('6.3  Audience Targeting')
    b.body(
        'Each package carries an audience object that defines who can see and use it. '
        'All rules are AND-combined — a user must satisfy every condition that is set. '
        'Rules that are omitted or null are ignored (not restrictive).'
    )

    b.table(
        headers=['Audience Rule', 'Data Source', 'Example Use Case'],
        rows=[
            ('allUsers: true',         '—',                        'Promo visible to everyone (overrides all other rules)'),
            ('roles: ["corporate"]',   'user.roles',               'Packages only for corporate account holders'),
            ('minOrderCount: 10',      'Count of completed orders', 'Loyalty rewards for frequent users'),
            ('activeWithinDays: 30',   'Most recent completed order date', 'Re-engagement offers for lapsed users'),
            ('areaIds: ["uuid-lekki"]','Customer\'s order area history', 'Localised promotions for specific neighbourhoods'),
            ('companyIds: ["uuid-X"]', 'CompanyEmployee membership', 'Corporate-exclusive packages for employees of a specific company'),
            ('requirePhone: true',     'user.phone',               'Packages only available to verified phone users'),
            ('requireAddress: true',   'Count of saved addresses',  'Packages that require a delivery address'),
        ],
        col_widths=[1.8, 1.8, 2.7]
    )

    b.h3('How audience filtering works at runtime:')
    b.numbered('Customer hits GET /pricing/packages (authenticated).')
    b.numbered('All active, non-expired packages are loaded from the database.')
    b.numbered('For any package with allUsers: true, it is immediately included.')
    b.numbered('For others, the system builds an AudienceContext for the user: '
               'fetches user profile, address count, completed order count and dates, '
               'order area IDs, and company memberships — all in one parallel query batch.')
    b.numbered('Each package\'s audience rules are evaluated against the context.')
    b.numbered('Only packages where all conditions pass are returned to the customer.')

    b.note(
        'The audience criteria system is designed to be extended. Future versions can add '
        'occupation-based targeting, loyalty tiers, referral status, or any other user attribute '
        'without changing the data model — just add new fields to the PackageAudience interface.',
        label='PRODUCT ROADMAP NOTE', color=BRAND_GREEN
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — VENDOR PRICE INTELLIGENCE ENGINE
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('7.  Vendor Price Intelligence Engine')

    b.lead(
        'The Intelligence Engine analyses what laundry vendors on the platform charge '
        'their customers for each type of garment, and uses that data to suggest '
        'optimal platform prices — ensuring the platform stays competitive while '
        'protecting its margin above the market median.'
    )

    b.h2('7.1  The Problem It Solves')
    b.body(
        'Platform prices (bag prices, special item prices) are set by admin in the '
        'price catalogue. Without data, these are arbitrary guesses. The Intelligence Engine '
        'replaces guesswork with market evidence: if 50 vendors on the platform charge '
        'between ₦600 and ₦1,200 to wash a suit, the platform should price itself '
        'above the midpoint to protect margin — but not so far above that customers '
        'choose individual vendors over the platform.'
    )

    b.h2('7.2  How Data Is Collected')
    b.body(
        'Each Vendor has a VendorPricing record containing their price list as an array '
        'of { garmentType, priceNaira } objects. Pricing updates are append-only; '
        'the most recent approved row per vendor is the "active" price. '
        'Vendors set prices in Naira; the system converts to WP for platform pricing.'
    )

    b.h2('7.3  Statistical Analysis')
    b.body('For each garment type that appears in at least one vendor\'s active price list, the engine computes:')

    b.table(
        headers=['Statistic', 'Meaning', 'Use Case'],
        rows=[
            ('Vendor count',   'How many vendors price this garment type',      'Confidence indicator — ignore types with < 3 vendors'),
            ('Min',            'Cheapest vendor on the platform',               'Floor — platform should never price below this'),
            ('Mean',           'Average price across all vendors',              'Baseline reference'),
            ('P25',            '25th percentile — 25% of vendors charge less',  'Budget end of the market'),
            ('P50 (Median)',   '50th percentile — half above, half below',      'True market midpoint'),
            ('P70',            '70th percentile — platform default suggestion', 'Above most vendors; protects margin'),
            ('P75',            '75th percentile',                               'More aggressive margin protection'),
            ('P90',            '90th percentile — top of market',              'Premium positioning option'),
            ('Max',            'Most expensive vendor on the platform',         'Ceiling reference'),
        ],
        col_widths=[1.0, 2.4, 2.4]
    )

    b.h3('Why P70 as the default?')
    b.body(
        'Setting the platform at the 70th percentile means the platform price is higher '
        'than 70% of individual vendors but lower than 30%. '
        'This is deliberately above the median because the platform provides additional '
        'value: professional vetting of vendors, quality guarantees, contactless collection '
        'and delivery, dispute resolution, and loyalty rewards. '
        'The percentile is admin-configurable (priceSuggestionPercentile in PlatformConfig) '
        'so the positioning can be changed without code deployment.'
    )

    b.h2('7.4  From Naira to WashPoints')
    b.body(
        'Vendor prices are stored in Naira. The suggestion is converted to WP using the '
        'current active conversion rate: '
    )
    b.code('suggestedWP = round(suggestedNaira × pointsPerUnit)')
    b.body('For example: P70 = ₦900 Naira × 10 WP/₦1 = 9,000 WP')

    b.h2('7.5  Applying Suggestions')
    b.body(
        'The admin can review the full intelligence report at GET /pricing/intelligence '
        'and then apply suggestions via POST /pricing/intelligence/apply.'
    )

    b.h3('Apply options:')
    b.bullet('garmentTypes — apply only for specific garment types (leave empty for all)')
    b.bullet('toleranceWP — skip suggestions where the difference from the current price is ≤ this value (avoids noise updates for tiny changes)')

    b.body(
        'Each applied suggestion creates a new entry in the platform_price_list table '
        'with priceType = "special_item" and effectiveFrom = now(). '
        'The existing entries are not deleted — they become historical records.'
    )

    b.h3('Intelligence Report Example:')
    b.table(
        headers=['Garment Type', 'Vendors', 'Min ₦', 'P50 ₦', 'P70 ₦', 'Suggested WP', 'Current WP', 'Diff WP'],
        rows=[
            ('shirt',   '38',  '400',  '650',   '820',   '8,200',  '7,500',  '+700'),
            ('suit',    '29',  '800',  '1,100', '1,350', '13,500', '12,000', '+1,500'),
            ('agbada',  '21',  '1,200','1,800', '2,200', '22,000', '22,000', '0'),
            ('duvet',   '17',  '500',  '800',   '950',   '9,500',  '10,000', '-500'),
            ('babygrow','8',   '200',  '320',   '390',   '3,900',  '0',      '+3,900'),
        ],
        col_widths=[1.1, 0.7, 0.6, 0.6, 0.6, 1.1, 1.0, 0.7]
    )
    b.body('All values illustrative.', italic=True, color=BRAND_SLATE)

    b.note(
        'The intelligence engine is the platform\'s competitive moat. By continuously '
        'tracking vendor market prices, Washermann can ensure its pricing always reflects '
        'real market conditions — no manual research required.',
        label='INVESTOR NOTE', color=BRAND_GREEN
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — PLATFORM CONFIGURATION
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('8.  Platform Configuration')

    b.lead(
        'All financial parameters live in a single, admin-editable configuration row. '
        'Changes take effect immediately on the next order — no deployment required.'
    )

    b.table(
        headers=['Parameter', 'Type', 'Default', 'Effect When Changed'],
        rows=[
            ('platformPriceOffsetPercent', 'Decimal %', '25%',       'Baseline markup over vendor median — used as reference for Intelligence Engine'),
            ('serviceChargePercent',       'Decimal %', '5%',        'Immediately raises/lowers the fee charged on every new order subtotal'),
            ('vatPercent',                 'Decimal %', '0%',        'Enables/disables VAT on all new orders (0 = disabled)'),
            ('repSharePercent',            'Decimal %', '15%',       'Controls how much of the order total is allocated to Rep earnings pool'),
            ('payoutRateNairaPerWP',       'Decimal',   '₦9/WP',    'Changes the Naira value vendors/reps receive per WP when cashing out'),
            ('lowRatingThreshold',         'Decimal',   '3.5',       'Rating below which a Rep is auto-flagged for admin review'),
            ('bonusCyclePeriod',           'Enum',      'monthly',   'Period over which Rep bonus earnings are accumulated before payout calculation'),
            ('orderAutoCompleteHours',     'Integer',   '24 hours',  'Hours after delivery confirmation before order auto-completes if customer is silent'),
            ('priceSuggestionPercentile',  'Integer',   '70',        'Which percentile of vendor prices the Intelligence Engine suggests as platform price'),
        ],
        col_widths=[2.1, 0.7, 0.9, 2.6]
    )

    b.h2('8.1  Guard Rails')
    b.body('Every parameter has validation constraints to prevent accidental misconfiguration:')
    b.table(
        headers=['Parameter', 'Min', 'Max'],
        rows=[
            ('platformPriceOffsetPercent', '0%',  '200%'),
            ('serviceChargePercent',       '0%',  '30%'),
            ('vatPercent',                 '0%',  '30%'),
            ('repSharePercent',            '0%',  '50%'),
            ('lowRatingThreshold',         '1.0', '5.0'),
            ('orderAutoCompleteHours',     '1h',  '168h (1 week)'),
            ('priceSuggestionPercentile',  '50',  '95'),
        ],
        col_widths=[2.5, 1.5, 1.5]
    )

    b.h2('8.2  Who Can Change Config')
    b.body('Only users with the ADMIN role can update platform configuration. '
           'Finance role users can read it. All other roles have no access.')

    b.h2('8.3  Audit Trail for Config Changes')
    b.body('The updatedBy field on PlatformConfig stores the UUID of the admin who last '
           'made a change. The updatedAt timestamp records when. '
           'The config table also has a createdAt. '
           'For full audit purposes, future versions can wrap config updates in a '
           'ConfigChangeLog table (the architecture supports this without model changes).')

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — ORDER FINANCIAL FLOW
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('9.  Order Financial Flow (End-to-End)')

    b.lead(
        'This section traces the complete money trail from the moment a customer '
        'places an order to the moment the vendor receives their earnings and '
        'can request a payout.'
    )

    b.h2('9.1  Pre-Order Checks')
    b.body('Before any financial operation begins, two gates are checked:')
    b.numbered('Profile Completion Gate — the customer must have a phone number and at least one saved delivery address. If either is missing, the order is rejected with a clear error listing the missing fields.')
    b.numbered('Balance Check — the customer\'s WashPoint balance must be ≥ the authoritative totalWP calculated by the server. If insufficient, the order is rejected.')

    b.h2('9.2  Order Placement Transaction')
    b.body('All of the following happen inside a single atomic database transaction:')
    b.table(
        headers=['Step', 'What Happens', 'Database Records Created/Modified'],
        rows=[
            ('1', 'Profile and balance pre-checks (outside transaction)', 'None yet'),
            ('2', 'Authoritative pricing calculation via PricingEngine',  'None — pure computation'),
            ('3', 'Customer wallet debited by totalWP',                   'Wallet.balance decremented; Wallet.fiatBalanceKobo reduced proportionally (WACB)'),
            ('4', 'Debit LedgerEntry written',                           'LedgerEntry (type=debit, source=ORDER_DEBIT, amount, balanceBefore, balanceAfter, conversionRateId/Snapshot)'),
            ('5', 'Order record created',                                 'Order (all fields; pricingSnapshot stores the full PricingResult JSON; status=PENDING_PAYMENT)'),
            ('6', 'Escrow record created',                                'OrderEscrow (holdWP = totalWP; locked until vendor completes; status=HELD)'),
            ('7', 'Order status history entry',                           'OrderStatusHistory (from=null, to=PENDING_PAYMENT)'),
            ('8', 'Assignment broadcast created',                         'AssignmentBroadcast (notifies available reps in the area)'),
        ],
        col_widths=[0.3, 2.8, 3.2]
    )

    b.h2('9.3  Order Lifecycle & Earnings Release')
    b.body('Earnings are not released immediately. They follow the order through its lifecycle:')
    b.table(
        headers=['Status', 'Meaning', 'Financial Action'],
        rows=[
            ('PENDING_PAYMENT',   'Order placed, awaiting Rep assignment',         'Funds in escrow'),
            ('ACCEPTED',          'Rep accepted the broadcast',                    'Funds in escrow'),
            ('PICKED_UP',         'Rep has collected garments from customer',      'Earnings split calculated and locked into Order (vendorShareWP, repShareWP, platformFeeWP)'),
            ('AT_VENDOR',         'Rep delivered to vendor; vendor processing',    'Funds in escrow'),
            ('READY_FOR_PICKUP',  'Vendor finished; ready for rep to collect',     'Funds in escrow'),
            ('DELIVERED',         'Rep returned clean laundry to customer',        'Funds in escrow'),
            ('COMPLETED',         'Customer confirmed, or 24h auto-complete fired','Escrow released: vendor credited, rep pseudo-wallet credited, platform retains remainder'),
            ('CANCELLED',         'Order cancelled (before PICKED_UP)',             'Full refund to customer wallet; fiatBalanceKobo restored proportionally'),
        ],
        col_widths=[1.5, 2.0, 2.8]
    )

    b.h2('9.4  Earnings Split Formula')
    b.body('At PICKED_UP status, the order\'s earnings are locked:')
    b.code(
        'vendorShareWP  = round(totalWP × vendorSharePercent / 100)\n'
        'repShareWP     = round(totalWP × repSharePercent / 100)\n'
        'platformFeeWP  = totalWP - vendorShareWP - repShareWP\n\n'
        '// vendorSharePercent = 100 - repSharePercent - platformFeePercent\n'
        '// These values are snapshotted at the time of PICKED_UP, not at order completion.'
    )

    b.note(
        'Locking the split at PICKED_UP (rather than COMPLETED) prevents the platform from '
        'changing fee structures between a service being rendered and payment being released. '
        'This is important for regulatory compliance and vendor trust.',
        label='COMPLIANCE NOTE', color=BRAND_AMBER
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — ESCROW SYSTEM
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('10.  Escrow System')

    b.lead(
        'Every order\'s WashPoints are held in escrow from the moment of placement '
        'until the order is confirmed complete. This protects customers from paying '
        'for services not rendered, while assuring vendors they will be paid.'
    )

    b.h2('10.1  How Escrow Works')
    b.body('The OrderEscrow entity holds:')
    b.bullet('holdWP — the exact WP amount held (= order totalWP)')
    b.bullet('status — HELD, RELEASED, or REFUNDED')
    b.bullet('releasedAt — timestamp when released to vendor')
    b.bullet('refundedAt — timestamp if order was cancelled and funds returned')

    b.h2('10.2  Escrow Release Conditions')
    b.table(
        headers=['Condition', 'Action', 'Who Triggers'],
        rows=[
            ('Customer confirms order complete',               'Escrow RELEASED → vendor credited, rep credited',         'Customer (explicit confirmation)'),
            ('24 hours pass after DELIVERED status',           'Auto-complete → escrow RELEASED → same as above',         'System (scheduled job)'),
            ('Order cancelled before PICKED_UP',               'Escrow REFUNDED → full WP returned to customer wallet',   'Customer or Admin'),
            ('Order cancelled after PICKED_UP (exceptional)',  'Partial or full refund depending on admin decision',      'Admin only'),
        ],
        col_widths=[2.4, 2.0, 1.8]
    )

    b.h2('10.3  Why Escrow Matters (Investor Perspective)')
    b.body(
        'Escrow is foundational to marketplace trust. Neither party can be defrauded: '
        'the customer cannot place an order they cannot pay for (balance checked first), '
        'and the vendor cannot be skipped on payment for work they have done. '
        'This trust mechanism is a key differentiator that enables Washermann to '
        'onboard professional vendors who would otherwise demand cash upfront.'
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — CUSTOMER WALLET
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('11.  Customer Wallet & Fiat Cost-Basis Tracking')

    b.h2('11.1  Wallet Structure')
    b.table(
        headers=['Field', 'Type', 'Purpose'],
        rows=[
            ('balance',          'bigint (WP)',  'Current spendable WashPoint balance'),
            ('fiatBalanceKobo',  'bigint (kobo)', 'Running Naira cost basis of the WP balance (WACB method)'),
            ('userId',           'uuid',          'One wallet per user'),
        ],
        col_widths=[1.8, 1.3, 3.2]
    )

    b.h2('11.2  LedgerEntry — The Audit Trail')
    b.body('Every wallet change writes an immutable LedgerEntry:')
    b.table(
        headers=['Field', 'Purpose'],
        rows=[
            ('type',                    '"credit" or "debit"'),
            ('amount',                  'WP moved (always positive)'),
            ('balanceBefore',           'WP balance before this event'),
            ('balanceAfter',            'WP balance after this event'),
            ('source',                  'Why the movement happened (TOP_UP, ORDER_DEBIT, REFUND, GIFT_CARD_REDEMPTION, etc.)'),
            ('conversionRateId',        'Rate active at this moment (for TOP_UP entries)'),
            ('conversionRateSnapshot',  'pointsPerUnit frozen at the time'),
            ('fiatAmountKobo',          'Naira paid (for TOP_UP entries only)'),
            ('reference',               'Order reference or Paystack reference — traceable to the source event'),
        ],
        col_widths=[2.0, 4.3]
    )

    b.h2('11.3  LedgerSource Values')
    b.body('The source field uses a controlled enum. Current values:')
    b.table(
        headers=['Source', 'When Written'],
        rows=[
            ('TOP_UP',               'Paystack payment completed and WP credited'),
            ('ORDER_DEBIT',          'Customer places an order — WP debited'),
            ('ORDER_REFUND',         'Order cancelled — WP returned'),
            ('GIFT_CARD_REDEMPTION', 'Customer redeems a gift card code'),
            ('VAULT_CREDIT',         'Admin or system credits WP directly from a vault'),
            ('COMPANY_TRANSFER',     'Company wallet distributes WP to employee wallets'),
        ],
        col_widths=[2.0, 4.3]
    )

    b.h2('11.4  Why the Ledger Is Immutable')
    b.body(
        'No ledger entry is ever modified or deleted in the service layer. '
        'This is a deliberate architectural choice that provides:'
    )
    b.bullet('Complete, tamper-proof financial history for every account')
    b.bullet('Ability to replay all transactions and reconstruct any wallet balance at any point in time')
    b.bullet('Regulatory compliance — financial records must be preserved')
    b.bullet('Fraud detection — balance drift (wallet.balance ≠ sum of ledger entries) immediately flags a discrepancy')

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 12 — VENDOR EARNINGS SYSTEM
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('12.  Vendor Earnings System')

    b.lead(
        'Vendors earn WashPoints for every completed order. These accumulate in a '
        'dedicated VendorEarningsWallet separate from the main customer wallet. '
        'Vendors can request a Naira payout at any time.'
    )

    b.h2('12.1  VendorEarningsWallet Structure')
    b.table(
        headers=['Field', 'Purpose'],
        rows=[
            ('balance',      'Current spendable WP balance — decremented when payout requested'),
            ('totalEarned',  'All-time WP earned — never decremented (for lifetime stats)'),
            ('status',       '"active" or "frozen" — frozen wallets cannot request payouts'),
        ],
        col_widths=[1.5, 4.8]
    )

    b.h2('12.2  VendorLedgerEntry')
    b.body('Like the customer ledger, all vendor wallet movements are append-only:')
    b.bullet('VENDOR_EARNING — WP credited after order COMPLETED')
    b.bullet('PAYOUT_DEBIT — WP debited when a payout request is processed')
    b.body('Each entry also stores nairaSnapshot — the Naira equivalent at the time, for the vendor\'s own records.')

    b.h2('12.3  Vendor Pricing Data Feed')
    b.body(
        'Vendors submit their own price list for each garment type (in Naira) via the '
        'VendorPricing system. This data feeds the Intelligence Engine (Section 7). '
        'Vendor pricing is append-only and requires admin approval before taking effect.'
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 13 — REP PSEUDO-WALLET & BONUS TIERS
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('13.  Rep Pseudo-Wallet & Bonus Tiers')

    b.lead(
        'Reps (delivery personnel) earn WashPoints for each delivery they complete. '
        'These are tracked in a "pseudo-wallet" — an internal ops tool, not visible '
        'to the Rep — and converted to a Naira salary + bonus each cycle.'
    )

    b.h2('13.1  Why a Pseudo-Wallet?')
    b.body(
        'Reps are employees or contractors, not marketplace participants. They are paid '
        'a salary, not a withdrawal from a balance. The pseudo-wallet is an internal '
        'accounting tool that lets ops compute earnings precisely and audit them, '
        'while the Rep\'s external pay is a normal bank transfer.'
    )

    b.h2('13.2  RepPseudoWallet Fields')
    b.table(
        headers=['Field', 'Purpose'],
        rows=[
            ('balance',        'Current cycle WP — reset to 0 at the start of each new bonus cycle'),
            ('totalEarned',    'All-time WP earned — never reset'),
            ('cycleStartedAt', 'When the current cycle began — used to calculate cycle duration'),
        ],
        col_widths=[1.6, 4.7]
    )

    b.h2('13.3  Bonus Tier System')
    b.body(
        'At the end of each cycle, each Rep\'s average rating determines their bonus percentage. '
        'The bonus is applied to their cycle WP earnings before Naira conversion.'
    )
    b.table(
        headers=['Tier', 'Rating Band', 'Bonus %', 'Flag for Review?'],
        rows=[
            ('Elite',  '4.8 – 5.0', '+15%', 'No'),
            ('Gold',   '4.5 – 4.7', '+10%', 'No'),
            ('Silver', '4.0 – 4.4', '+5%',  'No'),
            ('Bronze', '3.5 – 3.9', '0%',   'No'),
            ('Review', '0.0 – 3.4', '0%',   ('Yes — sent to admin review', False, True)),
        ],
        col_widths=[1.0, 1.2, 0.9, 3.2]
    )

    b.body('Bonus tiers are fully admin-configurable via the upsertBonusTier endpoint. '
           'Multiple tiers can be active simultaneously; the system picks the one matching '
           'the rep\'s rating at cycle end.')

    b.h2('13.4  Bonus Calculation Formula')
    b.code(
        'cyclePayNaira = (cycleWP × payoutRateNairaPerWP) × (1 + bonusPercent / 100)\n\n'
        '// Example: Rep earned 5,000 WP in the "Gold" tier (10% bonus)\n'
        '// payoutRate = ₦9/WP\n'
        '// cyclePayNaira = (5,000 × 9) × 1.10 = ₦45,000 × 1.10 = ₦49,500'
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 14 — PAYOUT SYSTEM
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('14.  Payout System')

    b.lead(
        'Vendors can request a Naira payout of their accumulated WashPoint earnings '
        'at any time. Payouts are processed via Paystack transfers to the vendor\'s '
        'registered bank account.'
    )

    b.h2('14.1  Payout Request Flow')
    b.numbered('Vendor submits a payout request (amountWP, bankCode, accountNumber, accountName).')
    b.numbered('System calculates nairaAmount = amountWP × payoutRateNairaPerWP. '
               'The current payout rate is snapshotted into the request — rate changes after submission do not affect pending requests.')
    b.numbered('VendorEarningsWallet balance is debited immediately.')
    b.numbered('Request enters PENDING status.')
    b.numbered('Admin reviews and approves.')
    b.numbered('System initiates Paystack transfer.')
    b.numbered('Paystack webhook updates request to COMPLETED or FAILED.')
    b.numbered('On FAILED, the WP are returned to the vendor\'s earnings wallet.')

    b.h2('14.2  Payout Request Fields')
    b.table(
        headers=['Field', 'Purpose'],
        rows=[
            ('amountWP',              'WashPoints being cashed out'),
            ('nairaAmount',           'Naira to be transferred (locked at request time)'),
            ('payoutRateSnapshot',    'Rate used for this request (never changes after creation)'),
            ('bankCode / accountNumber / accountName', 'Destination bank details'),
            ('status',                'PENDING → PROCESSING → COMPLETED or FAILED'),
            ('paystackReference',     'Paystack transfer reference for reconciliation'),
            ('batchId',               'Optional — links this request to a bulk payout run'),
            ('approvedBy / approvedAt', 'Admin who approved and when'),
        ],
        col_widths=[2.5, 3.8]
    )

    b.note(
        'The payout rate (payoutRateNairaPerWP) is intentionally separate from the customer '
        'buy rate (pointsPerUnit on ConversionRate). The buy rate determines how many WP '
        'customers get per Naira paid in. The payout rate determines how many Naira vendors '
        'receive per WP cashed out. The spread is platform margin.',
        label='INVESTOR NOTE', color=BRAND_GREEN
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 15 — GIFT CARDS & VAULTS
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('15.  Gift Cards & Vaults')

    b.lead(
        'Gift Cards allow WashPoints to be distributed to customers through promotional '
        'codes. Vaults are the pools of WashPoints from which gift cards are funded.'
    )

    b.h2('15.1  Vaults')
    b.body(
        'A Vault is a pre-funded pool of WashPoints created by admin. Think of it as '
        'a budget allocation for a specific promotional campaign or purpose.'
    )
    b.table(
        headers=['Vault Field', 'Purpose'],
        rows=[
            ('purpose',              'GENERAL, GIFT_CARDS, COUPONS, or CUSTOM — organises vaults by campaign type'),
            ('totalPoints',          'Total WP the vault can issue — fixed at creation'),
            ('usedPoints',           'WP issued so far — incremented on each debit'),
            ('conversionRateSnapshot', 'Rate locked at vault creation — all gift cards from this vault use this rate'),
            ('sequenceOrder',        'When multiple vaults exist, lower number = higher priority for auto-selection'),
            ('autoCreateOnThreshold','When usedPoints approaches totalPoints, automatically create the next vault'),
            ('status',               'ACTIVE, EXHAUSTED (no points left), or DEACTIVATED'),
        ],
        col_widths=[2.0, 4.3]
    )

    b.h2('15.2  Gift Cards')
    b.body(
        'Gift cards are single-use (or multi-use) redemption codes generated from a vault. '
        'When a customer redeems a gift card, the specified WP are credited to their wallet '
        'and debited from the source vault.'
    )
    b.table(
        headers=['Gift Card Feature', 'Detail'],
        rows=[
            ('code',                 '24-character unique redemption code'),
            ('wpValuePerUse',        'WP credited per redemption'),
            ('maxUsages',            'How many times the code can be used (default: 1)'),
            ('qualificationCriteria','JSON rules for who can redeem (e.g. employeeOnly, companyId)'),
            ('isPublic',             'Whether non-employees can use a company-issued card'),
            ('expiresAt',            'Optional expiry date'),
            ('status',               'ACTIVE, EXHAUSTED, REVOKED, EXPIRED'),
        ],
        col_widths=[2.0, 4.3]
    )

    b.h2('15.3  Company Gift Cards')
    b.body(
        'Companies can issue gift cards from their Company Wallet. This enables corporate '
        'benefit programs: a company tops up its corporate wallet and distributes gift card '
        'codes to employees, who redeem them for laundry orders.'
    )

    b.note(
        'Gift card codes are pre-debited from their source (vault or company wallet) at '
        'creation time, not at redemption time. This prevents the source from being '
        'over-committed if multiple redemptions happen simultaneously.',
        label='TECHNICAL NOTE', color=BRAND_GREEN
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 16 — AUDIT TRAIL & LEDGER IMMUTABILITY
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('16.  Audit Trail & Ledger Immutability')

    b.h2('16.1  The Three Ledgers')
    b.table(
        headers=['Ledger', 'Who It Belongs To', 'Purpose'],
        rows=[
            ('LedgerEntry',          'Customer wallets',  'All WP credits and debits for customers (top-ups, order payments, refunds, gift cards)'),
            ('VendorLedgerEntry',    'Vendor earnings wallets', 'All WP credits (earnings) and debits (payout requests) for vendors'),
            ('RepPseudoLedgerEntry', 'Rep pseudo-wallets', 'All WP credits (delivery earnings) for reps — cycle resets included'),
        ],
        col_widths=[1.8, 1.7, 2.8]
    )

    b.h2('16.2  Immutability Guarantee')
    b.body('All three ledgers share the same architectural pattern:')
    b.bullet('No UPDATE or DELETE operations exist in the service layer for ledger records')
    b.bullet('Every state change appends a NEW record with balanceBefore and balanceAfter')
    b.bullet('The running balance on the wallet entity is a derived value — it should always equal the sum of all ledger entries (this is the reconciliation check)')
    b.bullet('All ledger entities have a CreateDateColumn but NO UpdateDateColumn — there is no update path')

    b.h2('16.3  Reconciliation')
    b.body('To verify wallet integrity:')
    b.code(
        'SELECT SUM(CASE WHEN type = \'credit\' THEN amount ELSE -amount END)\n'
        'FROM   ledger_entries\n'
        'WHERE  wallet_id = :walletId\n\n'
        '-- Result must equal wallet.balance\n'
        '-- Any discrepancy indicates data corruption or direct DB manipulation'
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 17 — PROFILE COMPLETION GATE
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('17.  Profile Completion Gate')

    b.body(
        'Before a customer can place an order, the system checks that their profile '
        'is sufficiently complete. An incomplete profile is a common source of order '
        'failures — if we allow orders without a phone number, delivery cannot be '
        'coordinated; without an address, the Rep has nowhere to collect from.'
    )

    b.h2('17.1  Required Fields')
    b.table(
        headers=['Requirement', 'Check', 'Why It Matters'],
        rows=[
            ('Phone number', 'user.phone is not null/empty',     'Rep calls customer to coordinate pickup; SMS notifications require a phone'),
            ('Saved address', 'At least 1 Address record exists', 'Rep needs a physical location to collect from; required for area routing'),
        ],
        col_widths=[1.4, 2.2, 2.7]
    )

    b.h2('17.2  What Happens on Failure')
    b.body('If either check fails, the system returns a structured 400 Bad Request:')
    b.code(
        '{\n'
        '  "message": "Profile incomplete — cannot place order",\n'
        '  "missingFields": ["phone", "address"],\n'
        '  "hint": "Add a phone number and at least one delivery address to your profile."\n'
        '}'
    )

    b.h2('17.3  Profile Completion Endpoint')
    b.body('GET /users/me/profile-completion returns a checklist:')
    b.code(
        '{\n'
        '  "isComplete": false,\n'
        '  "checks": { "phone": false, "address": true },\n'
        '  "missingFields": ["phone"],\n'
        '  "message": "Complete your profile to place orders. Missing: phone"\n'
        '}'
    )
    b.body('The app should call this endpoint on profile screens and display nudges to guide users to completion.')

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 18 — API ENDPOINT REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('18.  API Endpoint Reference')

    b.h2('18.1  Pricing Endpoints')
    b.table(
        headers=['Method', 'Path', 'Auth', 'Description'],
        rows=[
            ('POST',  '/pricing/calculate',              'Any',           'Authoritative price calculation for a specific order config'),
            ('GET',   '/pricing/model/:areaId',          'None',          'Full all-in pricing model for an area; cache ~5 min'),
            ('GET',   '/pricing/config/:areaId',         'None',          'Legacy alias for /pricing/model/:areaId'),
            ('GET',   '/pricing/packages',               'JWT required',  'List pricing packages visible to the authenticated user'),
            ('GET',   '/pricing/intelligence',           'Admin/Finance', 'Full vendor price intelligence report'),
            ('POST',  '/pricing/intelligence/apply',     'Admin only',    'Apply intelligence suggestions to platform price list'),
            ('POST',  '/pricing/packages/admin',         'Admin only',    'Create a new pricing package'),
            ('GET',   '/pricing/packages/admin',         'Admin/Finance', 'List all packages (including inactive)'),
            ('GET',   '/pricing/packages/admin/:id',     'Admin/Finance', 'Get a single package'),
            ('PATCH', '/pricing/packages/admin/:id',     'Admin only',    'Update a package'),
            ('DELETE','/pricing/packages/admin/:id',     'Admin only',    'Permanently delete a package'),
        ],
        col_widths=[0.7, 2.5, 1.2, 2.0]
    )

    b.h2('18.2  Platform Config Endpoints')
    b.table(
        headers=['Method', 'Path', 'Auth', 'Description'],
        rows=[
            ('GET',   '/platform-config',              'Admin/Finance', 'Get current platform configuration'),
            ('PATCH', '/platform-config',              'Admin only',    'Update configuration fields (partial update)'),
            ('POST',  '/platform-config/prices',       'Admin only',    'Add a new price list entry (bag, special_item, ironing)'),
            ('GET',   '/platform-config/prices',       'Admin/Finance', 'List all price entries (newest first)'),
            ('GET',   '/platform-config/bonus-tiers',  'Admin/Finance', 'List active Rep bonus tiers'),
            ('POST',  '/platform-config/bonus-tiers',  'Admin only',    'Create or update a bonus tier'),
            ('DELETE','/platform-config/bonus-tiers/:id','Admin only',  'Deactivate a bonus tier'),
        ],
        col_widths=[0.7, 2.6, 1.2, 2.0]
    )

    b.h2('18.3  User / Profile Endpoints (Financial Relevance)')
    b.table(
        headers=['Method', 'Path', 'Auth', 'Description'],
        rows=[
            ('GET',   '/users/me/profile-completion', 'JWT required', 'Check if profile is complete for order placement'),
            ('PATCH', '/users/me',                    'JWT required', 'Update phone number and other profile fields'),
            ('POST',  '/users/me/addresses',          'JWT required', 'Add a delivery address (required before first order)'),
        ],
        col_widths=[0.7, 2.5, 1.2, 2.0]
    )

    b.h2('18.4  Order Endpoints (Financial)')
    b.table(
        headers=['Method', 'Path', 'Auth', 'Description'],
        rows=[
            ('POST',  '/orders',         'JWT required',  'Place an order (profile gate + balance check + pricing + escrow creation)'),
            ('PATCH', '/orders/:id/confirm', 'JWT required', 'Customer confirms order complete → releases escrow'),
        ],
        col_widths=[0.7, 2.0, 1.2, 2.5]
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 19 — QA TEST SCENARIOS
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('19.  QA Test Scenarios')

    b.lead(
        'The following scenarios should be verified for every release. '
        'Each scenario includes the expected outcome and the database state to verify.'
    )

    b.h2('19.1  Pricing Engine Tests')
    b.table(
        headers=['Scenario', 'Input', 'Expected Result'],
        rows=[
            ('Basic bag order',
             'Medium Wash & Fold, no special items, 5% SC, 0% VAT, ₦150 transport',
             'subtotalWP = bag price; serviceChargeWP = 5%; vatWP = 0; totalWP = subtotal+SC+transport'),
            ('VAT enabled',
             'Same as above but vatPercent = 7.5',
             'vatWP = (subtotal + SC) × 7.5%; totalWP includes VAT'),
            ('Special items',
             'Small bag + 2 suits + 1 agbada',
             'Line items for bag + each item; subtotal = sum of all; fees applied on subtotal'),
            ('Wash & Iron',
             'Large bag, wash_iron, ironingCount = 5',
             'Ironing line item present at unitPriceWP × 5; subtotal includes ironing'),
            ('Zero ironing on wash_fold',
             'Medium bag, wash_fold, ironingCount = 10',
             'Ironing line item NOT present (ironingCount ignored for wash_fold)'),
            ('Transport area = 0',
             'Area with transportFeeWP = 0',
             'No transport line item; totalWP = subtotal + SC + VAT'),
            ('Unknown special item',
             'Item type not in price catalogue',
             'Item skipped (unitPrice = 0); no line item added'),
        ],
        col_widths=[1.8, 2.2, 2.3]
    )

    b.h2('19.2  Order Placement Tests')
    b.table(
        headers=['Scenario', 'Setup', 'Expected Outcome'],
        rows=[
            ('Insufficient balance',
             'User wallet = 500 WP; order total = 1,000 WP',
             '400 Bad Request: "Insufficient WashPoints"'),
            ('Missing phone (profile gate)',
             'User has no phone, has address',
             '400 Bad Request: missingFields = ["phone"]'),
            ('Missing address (profile gate)',
             'User has phone, no address',
             '400 Bad Request: missingFields = ["address"]'),
            ('Both missing',
             'User has no phone and no address',
             '400 Bad Request: missingFields = ["phone", "address"]'),
            ('Successful order',
             'Profile complete, sufficient balance',
             'Order created; wallet debited; escrow created; ledger entry written; fiatBalanceKobo reduced proportionally'),
            ('Correct fiatBalanceKobo deduction',
             'Wallet: 10,000 WP, fiatBalance = 1,000,000 kobo; order = 2,000 WP',
             'fiatBalance reduced by 20% → 800,000 kobo'),
        ],
        col_widths=[1.8, 1.9, 2.6]
    )

    b.h2('19.3  Intelligence Engine Tests')
    b.table(
        headers=['Scenario', 'Setup', 'Expected Result'],
        rows=[
            ('P70 calculation',
             '10 vendors pricing suit: ₦600, ₦700, ₦750, ₦800, ₦850, ₦900, ₦950, ₦1000, ₦1100, ₦1200',
             'P70 = ₦958.50 (linear interpolation on sorted array)'),
            ('Only active pricing used',
             'Vendor A has 2 pricing rows: old (expired) and new (active)',
             'Only new row\'s items are included in analysis'),
            ('Conversion to WP',
             'suggestedNaira = ₦900, pointsPerUnit = 10',
             'suggestedWP = 9,000'),
            ('Apply with tolerance',
             'Current platform price = 9,000 WP, suggested = 9,300 WP, toleranceWP = 500',
             'Skipped (diff = 300 ≤ tolerance 500)'),
            ('Apply creates new entry',
             'Current = 7,500 WP, suggested = 9,000 WP, toleranceWP = 0',
             'New platform_price_list entry created; old entry NOT deleted'),
        ],
        col_widths=[1.8, 2.2, 2.3]
    )

    b.h2('19.4  Pricing Package Audience Tests')
    b.table(
        headers=['Scenario', 'Package Audience', 'User State', 'Visible?'],
        rows=[
            ('allUsers package',       '{ allUsers: true }',                   'Any user',                      'Yes'),
            ('minOrderCount = 5',      '{ minOrderCount: 5 }',                 'User has 4 completed orders',   'No'),
            ('minOrderCount = 5',      '{ minOrderCount: 5 }',                 'User has 5 completed orders',   'Yes'),
            ('Role check: corporate',  '{ roles: ["corporate"] }',             'User role = ["user"]',          'No'),
            ('activeWithinDays = 30',  '{ activeWithinDays: 30 }',             'Last order was 45 days ago',    'No'),
            ('requirePhone',           '{ requirePhone: true }',               'User has no phone',             'No'),
            ('areaIds match',          '{ areaIds: ["uuid-lekki"] }',         'User\'s orders are in Lekki area','Yes'),
            ('Package expired',        'validUntil = yesterday',               'Any user',                      'No'),
            ('Package not yet live',   'validFrom = tomorrow',                 'Any user',                      'No'),
            ('Package inactive',       'isActive = false',                     'Any user',                      'No'),
        ],
        col_widths=[1.8, 1.8, 1.7, 0.8]
    )

    b.h2('19.5  Wallet & Ledger Integrity Tests')
    b.table(
        headers=['Test', 'Verification Method'],
        rows=[
            ('Wallet balance = sum of ledger entries',
             'SELECT SUM(CASE WHEN type=\'credit\' THEN amount ELSE -amount END) FROM ledger_entries WHERE wallet_id = X'),
            ('No ledger entry is ever updated',
             'Check that no UpdateDateColumn exists on ledger entities; verify all DB rows have updated_at = created_at'),
            ('fiatBalanceKobo never goes negative',
             'After every order, assert wallet.fiatBalanceKobo >= 0'),
            ('Payout rate snapshot frozen',
             'Change payoutRateNairaPerWP after creating a payout request; verify request.payoutRateSnapshot unchanged'),
            ('Conversion rate snapshot frozen on transaction',
             'Change pointsPerUnit after initiating a Paystack transaction; verify transaction.conversionRateSnapshot unchanged'),
        ],
        col_widths=[2.3, 4.0]
    )

    b.page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 20 — GLOSSARY
    # ══════════════════════════════════════════════════════════════════════════
    b.h1('20.  Glossary')

    terms = [
        ('WashPoints (WP)',         'The internal currency of the Washermann platform. All prices, earnings, and transactions are denominated in WP.'),
        ('Conversion Rate',         'The official exchange rate between Naira and WashPoints. Set by admin; multiple rates can coexist with different effective dates.'),
        ('pointsPerUnit',           'WP per ₦1 — how many WashPoints a customer receives for each Naira spent.'),
        ('Buy Rate',                'The customer-facing conversion rate (how much WP per Naira paid).'),
        ('Payout Rate',             'The vendor/rep-facing rate (how much Naira per WP cashed out). Separate from buy rate.'),
        ('Rate Spread',             'The difference between buy rate and payout rate — one component of platform margin.'),
        ('PricingEngine',           'Stateless calculation class that converts inputs + config into a full PricingResult. No database access.'),
        ('PricingConfig',           'The set of prices and percentages the engine needs — loaded from DB before each calculation.'),
        ('PricingResult',           'The output of a PricingEngine.calculate() call — full itemised breakdown + totals.'),
        ('lineItems',               'Individual charge lines in a PricingResult (bag, special item, ironing, service charge, VAT, transport).'),
        ('Subtotal',                'Sum of item prices (bag + special items + ironing) before service charge, VAT, and transport.'),
        ('Service Charge',          'A percentage fee applied on the order subtotal. Default 5%.'),
        ('VAT',                     'Value Added Tax — configurable percentage, default 0 (disabled). Applied on subtotal + service charge.'),
        ('Transport Fee',           'Flat per-order fee set per Area. Covers the cost of Rep logistics.'),
        ('Rep Share',               'Percentage of the order total allocated to the delivery Rep earnings pool. Default 15%.'),
        ('Escrow',                  'A financial hold on the customer\'s WP from order placement until order completion. Protects both parties.'),
        ('WACB',                    'Weighted Average Cost Basis — accounting method used to track the Naira cost basis of WP balances.'),
        ('fiatBalanceKobo',         'Field on Wallet tracking the real Naira value of the WP balance, updated proportionally on each debit.'),
        ('LedgerEntry',             'An immutable, append-only record of a single wallet transaction.'),
        ('LedgerSource',            'Controlled enum identifying why a ledger entry was created (e.g. TOP_UP, ORDER_DEBIT, REFUND).'),
        ('VendorEarningsWallet',    'Dedicated wallet for a vendor\'s accumulated WP earnings, separate from customer wallets.'),
        ('RepPseudoWallet',         'Internal ops tool tracking Rep WP earnings. Not visible to the Rep — used to compute Naira salary.'),
        ('Bonus Tier',              'Admin-configured rating band that determines a Rep\'s bonus percentage for the cycle.'),
        ('Bonus Cycle',             'The period over which Rep earnings accumulate before salary calculation. Default: monthly.'),
        ('Payout Request',          'Vendor\'s request to convert WP earnings to Naira bank transfer.'),
        ('Vault',                   'Admin-created pool of WashPoints used to fund gift cards and promotions.'),
        ('Gift Card',               'Redemption code backed by a Vault or Company Wallet; credits WP to the redeemer\'s wallet.'),
        ('Pricing Package',         'Promotional or niche bundle with a fixed WP price, structured criteria, and audience targeting rules.'),
        ('Audience',                'JSON targeting rules on a PricingPackage defining who can see and use it.'),
        ('Intelligence Engine',     'Service that aggregates vendor prices per garment type and computes statistical suggestions for platform pricing.'),
        ('P70 / Percentile',        'Statistical measure. P70 means 70% of vendors price below this value. Used as the default platform price suggestion.'),
        ('Platform Price List',     'Append-only catalogue of WashPoint prices for all item types. Never modified — only new rows added.'),
        ('Profile Completion Gate', 'Mandatory check before order placement — requires phone number and at least one saved address.'),
        ('Append-Only',             'Architectural pattern where records are never updated or deleted. State changes create new records.'),
    ]

    for term, definition in terms:
        p = b.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(term + ':  ')
        r1.bold = True
        r1.font.color.rgb = BRAND_BLUE
        r2 = p.add_run(definition)
        r2.font.color.rgb = BRAND_DARK
        r2.font.size = Pt(10)

    b.divider()
    b.divider()

    # Footer note
    footer_para = b.doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        'This document is confidential and intended for authorised recipients only.\n'
        'Washermann © 2026 — All Rights Reserved'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_SLATE
    run.italic = True

    # ── Save ──────────────────────────────────────────────────────────────────
    b.save(path)


if __name__ == '__main__':
    import os
    out = os.path.join(os.path.dirname(__file__), 'Washermann_Financial_Pricing_Documentation.docx')
    build(out)
